# ============================================================
# FINAL MASTER VERSION 🚒
# KOMPLETTE FINALE VERSION
# ============================================================
# ENTHÄLT:
# ✅ Original Tabellenbreiten
# ✅ Statistik-Tabelle
# ✅ bold_tracking
# ✅ Poseidon:
#    - zwingend 1 Kadermitglied
#    - Kadermitglied fett markiert
# ✅ Grossreinigung:
#    - 3 Personen
#    - 2 Bootsführer zwingend
# ✅ GUI Jahr -> ersetzt <<JJJJ>>
# ============================================================

import pandas as pd
import os
import glob
import random
from datetime import datetime
from collections import defaultdict
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document

MONTHS = ["Jan","Feb","Mär","Apr","Mai","Jun","Jul","Aug","Sep","Okt","Nov","Dez"]
DARKBLUE = RGBColor(0, 51, 102)
EXCLUDE = ["Kommandant","Kommandant Stv","Chef Ausbildung"]


# ============================================================
# HELPERS
# ============================================================

def split(v):
    if pd.isna(v):
        return []
    return [x.strip() for x in str(v).split(",")]

def has(p, k):
    return k in p["g"] or k in p["f"]

def excluded(f):
    return any(x in EXCLUDE for x in f)

def kader(p):
    return "Kader" in p["f"]

def is_special(p):
    return any(x in p["f"] for x in EXCLUDE)

def basis_person(p):
    erlaubt = {"Mannschaft", "Fahrer"}

    funktionen = set(p["f"])

    return len(funktionen) > 0 and funktionen.issubset(erlaubt)


# ============================================================
# PLANNER
# ============================================================

class Planner:

    def __init__(self, persons):
        self.h = {
            p["id"]: {
                "last": -10,
                "count": 0,
                "veh": []
            }
            for p in persons
        }

        # Standardziel = 2 Einsätze
        total_slots = 0

        for m in range(12):
            total_slots += 2 + 2 + 2 + 2  # TLF, AS, SEE, POSE

        avg = total_slots // len(persons)

        self.targets = {
            p["id"]: avg
            for p in persons
        }

    def ok(self, p, m):
        return self.h[p["id"]]["last"] != m

    def score(self, p, vehicle):

        # Hauptregel: gleiche Anzahl Einsätze
        target = self.targets[p["id"]]
        count = self.h[p["id"]]["count"]

        if count >= target:
            count_score = 1000 + count * 50
        else:
            count_score = count * 10

        # Fahrzeugrotation stärker gewichten
        veh_count = self.h[p["id"]]["veh"].count(vehicle)

        # stärkere Bestrafung für gleiche Fahrzeuge
        veh_penalty = veh_count * 30

        # Bonus wenn Fahrzeug noch nie gefahren wurde
        if veh_count == 0:
            veh_penalty -= 5

        # zusätzliche Strafe für direkt hintereinander
        if self.h[p["id"]]["veh"]:
            if self.h[p["id"]]["veh"][-1] == vehicle:
                veh_penalty += 40

        bonus = 0

        # SEE bevorzugen beim SEE-Sprinter
        if vehicle == "SEE" and has(p, "SEE"):
            bonus -= 2

        # Basisleute bevorzugt einsetzen
        if basis_person(p):
            bonus -= 2

        return count_score + veh_penalty + bonus

    def assign(self, candidates, m, vehicle, need=None, n=2):

        # aktuelle Verteilung prüfen
        counts = [self.h[p["id"]]["count"] for p in candidates]
        min_count = min(counts) if counts else 0

        available = [
            p for p in candidates
            if self.ok(p, m)

            # 🔥 Spezialpersonen dürfen bei Grossreinigung nicht blockieren
            and not (
                is_special(p)
                and self.h[p["id"]]["count"] >= 1
                and m != getattr(self, "gross_month", -1)
            )
        ]

        # Falls zu strikt → lockern
        if not available:
            available = [
                p for p in candidates
                if self.ok(p, m)
                and not (is_special(p) and self.h[p["id"]]["count"] >= 1)
            ]

        # nach Anzahl sortieren
        random.shuffle(available)

        available = sorted(
            available,
            key=lambda p: (
                self.h[p["id"]]["count"],
                self.score(p, vehicle) + random.uniform(0, 3)
            )
        )

        s = []

        # Pflichtqualifikationen zuerst
        if need:
            for cond in need:

                for p in available:
                    if cond(p) and p not in s:
                        s.append(p)
                        break

        # Rest auffüllen
        for p in available:

            if len(s) >= n:
                break

            if p in s:
                continue

            s.append(p)

        # Falls immer noch zu wenig -> Monatsregel ignorieren
        if len(s) < n:

            fallback = sorted(
                [
                    p for p in candidates
                    if not (is_special(p) and self.h[p["id"]]["count"] >= 1)
                ],
                key=lambda p: self.h[p["id"]]["count"]
            )

            for p in fallback:

                if len(s) >= n:
                    break

                if p in s:
                    continue

                s.append(p)

        # speichern
        for p in s:
            self.h[p["id"]]["last"] = m
            self.h[p["id"]]["count"] += 1
            self.h[p["id"]]["veh"].append(vehicle)

        return s


# ============================================================
# DOCX HELPERS
# ============================================================

def set_col_widths(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Layout fixieren (WICHTIG!)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    widths_cm = [1.0, 3.75, 3.75, 3.75, 3.75]

    # Gesamtbreite setzen
    total_width = sum(widths_cm)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(total_width * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Grid definieren (entscheidend!)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_cm:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w * 567)))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # Spaltenbreiten setzen
    for i, w in enumerate(widths_cm):
        table.columns[i].width = Cm(w)

    # zusätzlich JEDE Zelle fixieren
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            cell = row.cells[i]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(w * 567)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)


def set_stats_widths(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    # feste Tabelle
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    widths_cm = [4.0, 2.0, 10.0]   # Person | Anzahl | Probefahrtsdaten

    total_width = sum(widths_cm)

    # Gesamtbreite
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(total_width * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Grid setzen
    tblGrid = OxmlElement('w:tblGrid')

    for w in widths_cm:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w * 567)))
        tblGrid.append(gridCol)

    tbl.append(tblGrid)

    # Spaltenbreite
    for i, w in enumerate(widths_cm):
        table.columns[i].width = Cm(w)

    # Jede Zelle fixieren
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            cell = row.cells[i]

            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(w * 567)))
            tcW.set(qn('w:type'), 'dxa')

            tcPr.append(tcW)


def shade(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "D9D9D9")
    tcPr.append(shd)


def set_bottom_border(cell):

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement("w:tcBorders")

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")

    tcBorders.append(bottom)
    tcPr.append(tcBorders)


def set_font(cell, bold=False, color=None):

    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = "Aptos"
            r.font.size = Pt(8.5)
            r.bold = bold

            if color:
                r.font.color.rgb = color

    for r in cell.paragraphs[0].runs:
        r.font.name = "Aptos"
        r.font.size = Pt(8.5)
        r.bold = bold


def replace_year(doc, year):

    placeholder = "<<JJJJ>>"
    value = str(year)

    # normale Absätze
    for p in doc.paragraphs:
        if placeholder in p.text:
            full = p.text.replace(placeholder, value)

            for run in p.runs:
                run.text = ""

            if p.runs:
                p.runs[0].text = full
            else:
                p.add_run(full)

    # Tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:

                    if placeholder in p.text:
                        full = p.text.replace(placeholder, value)

                        for run in p.runs:
                            run.text = ""

                        if p.runs:
                            p.runs[0].text = full
                        else:
                            p.add_run(full)

def write(cell, entries, bold_counter):

    cell.vertical_alignment = 0

    # Zelle leeren
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    bold_names = []

    # ------------------------------------
    # 1. forced Kandidaten
    # ------------------------------------
    forced_indices = [
        i for i, x in enumerate(entries)
        if len(x) >= 4 and x[3]
    ]

    # ------------------------------------
    # 2. normale Kandidaten (keine Spezialtexte)
    # ------------------------------------
    normal_indices = [
        i for i, x in enumerate(entries)
        if not x[2]
    ]

    chosen_index = None

    # ------------------------------------
    # PRIORITÄT 1: forced
    # ------------------------------------
    if forced_indices:
        chosen_index = forced_indices[0]

    # ------------------------------------
    # PRIORITÄT 2: faire Verteilung
    # ------------------------------------
    elif normal_indices:

        min_count = min(
            bold_counter[entries[i][0]]
            for i in normal_indices
        )

        kandidaten = [
            i for i in normal_indices
            if bold_counter[entries[i][0]] == min_count
        ]

        chosen_index = random.choice(kandidaten)

    # ------------------------------------
    # FALLBACK (z.B. nur Texte)
    # ------------------------------------
    elif entries:
        chosen_index = 0

    # ------------------------------------
    # SCHREIBEN
    # ------------------------------------
    for i, item in enumerate(entries):

        if len(item) == 3:
            txt, col, sp = item

        elif len(item) == 4:
            txt, col, sp, _ = item

        else:
            txt, col, sp, _, _ = item

        p = cell.add_paragraph()

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

        r = p.add_run(txt)
        r.font.name = "Aptos"
        r.font.size = Pt(8.5)

        # 👉 EINZIGE Stelle wo Fett gesetzt wird
        if i == chosen_index and col != "red":
            r.bold = True
            bold_names.append(txt)
            bold_counter[txt] += 1

        if col == "red":
            r.font.color.rgb = RGBColor(255, 0, 0)

    return bold_names


# ============================================================
# DOCX EXPORT
# ============================================================

def gen_doc(template, out, plan, stats, year):

    doc = Document(template)
    replace_year(doc, year)

    t = None

    for para in doc.paragraphs:

        if "<<PROBEFAHRT_PLAN>>" in para.text:

            parent = para._element.getparent()
            idx = parent.index(para._element)

            t = doc.add_table(rows=1, cols=5)
            parent.insert(idx, t._element)
            parent.remove(para._element)

            break

    if t is None:
        raise ValueError("Platzhalter <<PROBEFAHRT_PLAN>> fehlt")

    set_col_widths(t)

    headers = ["", "TLF", "Sprinter AS / MS",
               "Sprinter SEE", "Poseidon"]

    bold_tracking = defaultdict(list)
    bold_counter = defaultdict(int)

    vehicle_map = {
        "TLF": "TLF",
        "AS": "Sprinter AS / MS",
        "SEE": "Sprinter SEE",
        "POSE": "Poseidon"
    }

    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h

        for p in c.paragraphs:
            p.alignment = 0   # links

            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1

            for r in p.runs:
                r.font.name = "Aptos"
                r.font.size = Pt(9.5)
                r.bold = True
                r.font.color.rgb = DARKBLUE

        set_bottom_border(c)

        # vertikal mittig
        tcPr = c._tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)

    # Höhe Headerzeile
    header_row = t.rows[0]
    header_row.height = Pt(20)
    header_row.height_rule = 1

    for m in range(12):

        row = t.add_row().cells
        table_row = t.rows[-1]

        # Monatshöhe fixieren (für saubere Vertikalzentrierung)
        table_row.height = Pt(1)
        table_row.height_rule = 0

        row[0].text = MONTHS[m]

        for p in row[0].paragraphs:
            p.alignment = 0   # links, 1 = zentriert

            for r in p.runs:
                r.font.name = "Aptos"
                r.font.size = Pt(9.5)
                r.bold = True
                r.font.color.rgb = DARKBLUE

        for p in row[0].paragraphs:
            p.alignment = 1

        tcPr = row[0]._tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "top")
        tcPr.append(vAlign)

        # Prüfen ob keine Probefahrt
        no_drive = (
            len(plan[m]["TLF"]) == 1 and
            plan[m]["TLF"][0][0] == "Keine Probefahrt"
        )

        if no_drive:

            table_row.height = Pt(36)
            table_row.height_rule = 0

            # TLF schreiben
            write(row[1], plan[m]["TLF"], bold_counter)

            # TLF + AS + SEE zusammenführen
            merged = row[1].merge(row[2])
            merged = merged.merge(row[3])

            # horizontal zentrieren
            for p in merged.paragraphs:
                p.alignment = 1

            # vertikal zentrieren
            tcPr = merged._tc.get_or_add_tcPr()
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(qn("w:val"), "center")
            tcPr.append(vAlign)

            # Poseidon separat
            pose = plan[m]["POSE"]

            if (
                len(pose) == 1 and
                pose[0][0] == "Keine Probefahrt"
            ):

                cell = row[4]
                cell.text = "Keine Probefahrt"

                p = cell.paragraphs[0]
                p.alignment = 0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

                r = p.runs[0]
                r.font.name = "Aptos"
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(255, 0, 0)

                tcPr = cell._tc.get_or_add_tcPr()

                for el in tcPr.findall(qn("w:vAlign")):
                    tcPr.remove(el)

                vAlign = OxmlElement("w:vAlign")
                vAlign.set(qn("w:val"), "center")
                tcPr.append(vAlign)

            else:
                write(row[4], pose, bold_counter)

        else:
            for i, k in enumerate(["TLF", "AS", "SEE", "POSE"]):

                # Spezialfall: nur Poseidon = Keine Probefahrt
                if k == "POSE" and \
                   len(plan[m]["POSE"]) == 1 and \
                   plan[m]["POSE"][0][0] == "Keine Probefahrt":

                    cell = row[i + 1]

                    # komplette Zelle leeren
                    for p0 in cell.paragraphs:
                        p0._element.getparent().remove(p0._element)

                    # neuen sauberen Absatz erzeugen
                    p = cell.add_paragraph()
                    r = p.add_run("Keine Probefahrt")

                    p.alignment = 0
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)

                    r = p.runs[0]
                    r.font.name = "Aptos"
                    r.font.size = Pt(8.5)
                    r.font.color.rgb = RGBColor(255, 0, 0)

                    tcPr = cell._tc.get_or_add_tcPr()

                    for el in tcPr.findall(qn("w:vAlign")):
                        tcPr.remove(el)

                    vAlign = OxmlElement("w:vAlign")
                    vAlign.set(qn("w:val"), "center")
                    tcPr.append(vAlign)

                else:
                    cell = row[i + 1]

                    bold_names = write(cell, plan[m][k], bold_counter)

                    # normale Einträge oben ausrichten
                    tcPr = cell._tc.get_or_add_tcPr()

                    for el in tcPr.findall(qn("w:vAlign")):
                        tcPr.remove(el)

                    vAlign = OxmlElement("w:vAlign")
                    vAlign.set(qn("w:val"), "top")
                    tcPr.append(vAlign)

                    for name in bold_names:
                        bold_tracking[name].append(
                            f"{vehicle_map[k]} ({MONTHS[m]})"
                        )

        if m % 2 == 0:
            for c in row:
                shade(c)

    # Statistik an Platzhalter <<Personenzuteilung>>
    insert_para = None

    for para in doc.paragraphs:
        if "<<Personenzuteilung>>" in para.text:
            insert_para = para
            break

    if insert_para:

        parent = insert_para._element.getparent()
        idx = parent.index(insert_para._element)

        t2 = doc.add_table(rows=1, cols=3)
        parent.insert(idx, t2._element)

        parent.remove(insert_para._element)

    else:
        # Fallback falls Platzhalter fehlt
        doc.add_paragraph()
        t2 = doc.add_table(rows=1, cols=3)
        
    set_stats_widths(t2)

    t2.rows[0].cells[0].text = "Person"
    t2.rows[0].cells[1].text = "Anzahl"
    t2.rows[0].cells[2].text = "Probefahrtsdaten"

    for c in t2.rows[0].cells:
        set_font(c, True)
        set_bottom_border(c)

    def stat_name_key(item):
        full = item[0].split()

        # erster Teil = Grad entfernen
        rest = full[1:]

        if len(rest) == 1:
            vorname = ""
            nachname = rest[0]
        else:
            vorname = " ".join(rest[:-1])
            nachname = rest[-1]

        return nachname.lower(), vorname.lower()


    sorted_stats = sorted(stats.items(), key=stat_name_key)

    for name, data in sorted_stats:

        r = t2.add_row().cells
        row = t2.rows[-1]

        row.height = Pt(12)
        row.height_rule = 1

        parts = name.split()[1:]   # Grad entfernen

        if len(parts) == 1:
            display = parts[0]
        else:
            vorname = " ".join(parts[:-1])
            nachname = parts[-1]
            display = f"{nachname} {vorname}"

        r[0].text = display
        r[1].text = str(data["count"])
        for c in [r[0], r[1]]:
            p0 = c.paragraphs[0]
            p0.paragraph_format.space_before = Pt(0)
            p0.paragraph_format.space_after = Pt(0)
            p0.paragraph_format.line_spacing = 1

        p = r[2].paragraphs[0]
        p.clear()

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = 0
        p.paragraph_format.line_spacing = Pt(9)

        for i, d in enumerate(data["dates"]):

            if i > 0:
                p.add_run(", ")

            run = p.add_run(d)
            run.font.name = "Aptos"
            run.font.size = Pt(8.5)

            if d in bold_tracking[name]:
                run.bold = True

        set_font(r[0])
        set_font(r[1])

        for run in r[2].paragraphs[0].runs:
            run.font.name = "Aptos"
            run.font.size = Pt(8.5)

        for cell in r:
            tcPr = cell._tc.get_or_add_tcPr()

            vAlign = OxmlElement("w:vAlign")
            vAlign.set(qn("w:val"), "center")
            tcPr.append(vAlign)

    # ------------------------------------------------
    # LEERZEILE MIT TRENNLINIE OBEN
    # ------------------------------------------------
    r = t2.add_row().cells

    for c in r:
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        set_bottom_border(c)

    # ------------------------------------------------
    # TOTALZEILE UNTERHALB DER LINIE
    # ------------------------------------------------
    total_personen = len(stats)

    # Summe Spalte Anzahl
    total_einsaetze = sum(d["count"] for d in stats.values())

    # Anzahl Probefahrten / Jahr
    aktive_monate = sum(
        1 for m in range(12)
        if not (
            len(plan[m]["TLF"]) == 1 and
            plan[m]["TLF"][0][0] == "Keine Probefahrt"
        )
    )

    # 👉 4 Fahrzeuge pro aktivem Monat
    jahresfahrten = aktive_monate * 4

    r = t2.add_row().cells
    row = t2.rows[-1]

    row.height = Pt(12)
    row.height_rule = 1

    r[0].text = f"{total_personen} Personen"
    r[1].text = str(total_einsaetze)
    r[2].text = f"{jahresfahrten} Probefahrten / Jahr"

    for c in r:
        set_font(c, True)

        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

    doc.save(out)


# ============================================================
# MAIN
# ============================================================

def optimize_fairness(plan, persons, stats):

    by_name = {p["name"]: p for p in persons}

    def has_required(vehicle, names):

        persons_map = {p["name"]: p for p in persons}

        if vehicle == "TLF":
            return any(has(persons_map[n], "TLF Fahrer")
                       for n in names if n in persons_map)

        elif vehicle == "AS":
            return any(has(persons_map[n], "Maschinisten")
                       for n in names if n in persons_map)

        elif vehicle == "POSE":
            return any(has(persons_map[n], "Bootsführer")
                       for n in names if n in persons_map)

        return True  # SEE

    MAX_ROUNDS = 150

    for _ in range(MAX_ROUNDS):

        changed = False

        # sortiere Personen nach Einsätzen
        sorted_people = sorted(
            stats.items(),
            key=lambda x: x[1]["count"]
        )

        low_group = [
            n for n, d in sorted_people[:5]
            if n in by_name and not is_special(by_name[n]) and stats[n]["count"] > 0
        ]

        high_group = [
            n for n, d in sorted_people[-5:]
            if n in by_name and not is_special(by_name[n]) and stats[n]["count"] > 0
        ]

        best_move = None
        best_improvement = 0

        for hi in high_group:
            for lo in low_group:

                if hi == lo:
                    continue

                diff = stats[hi]["count"] - stats[lo]["count"]

                if diff <= 1:
                    continue

                # 🔥 niemals neue Ungleichheit erzeugen
                if stats[lo]["count"] + 1 > stats[hi]["count"] - 1:
                    continue

                for m in range(12):
                    for veh in ["TLF", "AS", "SEE", "POSE"]:

                        entries = plan[m][veh]

                        for i, item in enumerate(entries):

                            name = item[0]

                            # ❗ Schutz
                            if (len(item) >= 5 and item[4]) or any(x in name for x in EXCLUDE):
                                continue

                            if name != hi:
                                continue

                            lp = by_name.get(lo)
                            # ❗ KEINE Spezialpersonen rein tauschen
                            if is_special(lp):
                                continue
                            
                            if not lp:
                                continue

                            if is_special(lp):
                                continue

                            # Qualifikation prüfen
                            ok = False
                            if veh == "TLF":
                                ok = has(lp, "TLF Fahrer")
                            elif veh == "AS":
                                ok = has(lp, "Maschinisten")
                            elif veh == "SEE":
                                ok = True
                            elif veh == "POSE":
                                ok = has(lp, "Bootsführer")

                            if not ok:
                                continue

                            # Regelcheck
                            current_names = [x[0] for x in entries]
                            temp_names = current_names.copy()
                            temp_names[i] = lo

                            if not has_required(veh, temp_names):
                                continue

                            improvement = diff

                            if improvement > best_improvement:
                                best_improvement = improvement
                                best_move = (m, veh, i, hi, lo)

        # 👉 nichts mehr möglich
        if not best_move:
            continue

        # 👉 besten Move ausführen
        m, veh, i, hi, lo = best_move

        item = plan[m][veh][i]

        if len(item) >= 5:
            plan[m][veh][i] = (lo, item[1], item[2], False, False)
        else:
            plan[m][veh][i] = (lo, item[1], item[2])

        stats[hi]["count"] -= 1
        stats[lo]["count"] += 1

        changed = True

        if not changed:
            break

def rebuild_stats(plan, valid_names):

    stats = defaultdict(lambda: {
        "count": 0,
        "dates": []
    })

    mp = {
        "TLF": "TLF",
        "AS": "Sprinter AS / MS",
        "SEE": "Sprinter SEE",
        "POSE": "Poseidon"
    }

    for m in range(12):
        for veh in ["TLF", "AS", "SEE", "POSE"]:

            for item in plan[m][veh]:

                name = item[0]

                if name not in valid_names:
                    continue

                stats[name]["count"] += 1
                stats[name]["dates"].append(
                    f"{mp[veh]} ({MONTHS[m]})"
                )

    return stats

def build_plan(pl, persons, special_persons, disabled, pose_cfg,
               name_to_person, valid_names):

    plan = []
    stats = defaultdict(lambda: {"count": 0, "dates": []})
    special_pool = special_persons.copy()
    random.shuffle(special_pool)
    used_special_ids = set()

    gross_m = pose_cfg["gross_m"]
    gross_persons = [x for x in pose_cfg.get("gross_persons", []) if x]

    for m in range(12):

        e = {}

        # -------------------------------------------------
        # GROSSREINIGUNG
        # -------------------------------------------------
        if m == gross_m:

            selected = []

            # manuell gewählte Personen
            for name in gross_persons:
                p = name_to_person.get(name)
                if p:
                    selected.append(p)

            # 👉 Pool ohne Spezial + ohne bereits gewählte
            pool = [p for p in persons if not is_special(p) and p not in selected]

            # 👉 zuerst fehlende Bootsführer ergänzen
            random.shuffle(pool)

            for p in pool:
                if sum(1 for x in selected if has(x, "Bootsführer")) >= 2:
                    break
                if has(p, "Bootsführer"):
                    selected.append(p)

            # 👉 dann auf 3 auffüllen
            pool = [p for p in persons if not is_special(p) and p not in selected]
            random.shuffle(pool)

            for p in pool:
                if len(selected) >= 3:
                    break
                selected.append(p)

            for p in selected[:3]:
                pl.h[p["id"]]["last"] = m
                pl.h[p["id"]]["count"] += 1
                pl.h[p["id"]]["veh"].append("POSE")

                # 🔥 NEU
                if is_special(p):
                    used_special_ids.add(p["id"])

            names = []
            bold_done = False

            for p in selected[:3]:
                bold = False

                if has(p, "Bootsführer") and not bold_done:
                    bold = True
                    bold_done = True

                names.append((p["name"], None, True, bold, True))

            # ❗ Fallback falls KEIN Bootsführer drin ist
            if not bold_done and names:
                first = names[0]
                names[0] = (first[0], first[1], first[2], True, first[4])

            # Zusatztext für Grossreinigung
            extra = pose_cfg.get("gross_d", "").strip()

            if extra:
                names.append((f"Grossreinigung {extra}", "red", True, False))
            else:
                names.append(("Grossreinigung", "red", True, False))

            e["POSE"] = names
            

        # -------------------------------------------------
        # NORMALE FAHRZEUGE
        # -------------------------------------------------
        if m in disabled and m != gross_m:

            e["TLF"] = [("Keine Probefahrt", "red", True)]
            e["AS"] = [("Keine Probefahrt", "red", True)]
            e["SEE"] = [("Keine Probefahrt", "red", True)]

        else:

            e["TLF"] = [
                (p["name"], None, False, False)
                for p in pl.assign(
                    persons,
                    m, "TLF",
                    [lambda p: has(p, "TLF Fahrer")]
                )
            ]

            e["AS"] = [
                (p["name"], None, False, False)
                for p in pl.assign(
                    persons,
                    m, "AS",
                    [lambda p: has(p, "Maschinisten")]
                )
            ]

            e["SEE"] = [
                (p["name"], None, False, False)
                for p in pl.assign(
                    persons,
                    m, "SEE"
                )
            ]

        # -------------------------------------------------
        # POSEIDON
        # -------------------------------------------------
        if m == gross_m:
            pass  # bereits gesetzt durch Grossreinigung

        elif pose_cfg["disable"][m]:

            entries = [("Keine Probefahrt", "red", True)]

            if pose_cfg["text"][m]:
                entries.append((pose_cfg["text"][m], "red", True))

            e["POSE"] = entries

        else:

            pose = pl.assign(
                persons,
                m, "POSE",
                [
                    lambda p: has(p, "Bootsführer"),
                    lambda p: kader(p)
                ],
                2
            )

            names = []
            boots_done = False

            for p in pose:
                bold = False

                if has(p, "Bootsführer") and not boots_done:
                    bold = True
                    boots_done = True

                names.append((p["name"], None, True, bold, False))

            if pose_cfg["text"][m]:
                names.append((pose_cfg["text"][m], "red", True, False))

            e["POSE"] = names

        # SPEZIALPERSONEN (ERSATZLOGIK)
        if m not in disabled and m != gross_m:

            # nur Spezialpersonen ohne Einsatz
            available_special = [
                p for p in special_pool
                if pl.h[p["id"]]["count"] == 0
            ]

            if not available_special:
                pass
            else:
                sp = available_special.pop(0)
                special_pool.remove(sp)  # sauber entfernen

                vehicles = ["TLF", "AS", "SEE", "POSE"]
                random.shuffle(vehicles)

                for veh in vehicles:

                    entries = e[veh]

                    for i, item in enumerate(entries):

                        name = item[0]
                        p = name_to_person.get(name)

                        if not p:
                            continue

                        # Qualifikation prüfen
                        if veh == "TLF" and not has(sp, "TLF Fahrer"):
                            continue
                        if veh == "AS" and not has(sp, "Maschinisten"):
                            continue
                        if veh == "POSE" and not has(sp, "Bootsführer"):
                            continue

                        # ersetzen
                        e[veh][i] = (sp["name"], None, False)

                        pl.h[sp["id"]]["count"] += 1
                        pl.h[sp["id"]]["last"] = m
                        pl.h[sp["id"]]["veh"].append(veh)

                        break
                    else:
                        continue
                    break

        # -------------------------------------------------
        # STATISTIK
        # -------------------------------------------------
        for k in e:
            for item in e[k]:

                n = item[0]

                if n not in valid_names:
                    continue

                stats[n]["count"] += 1
                stats[n]["dates"].append(f"{k} ({MONTHS[m]})")

        plan.append(e)

    return plan, stats

def run(file, template, year, disabled, pose_cfg):

    df = pd.read_excel(file)
    gross_persons = pose_cfg.get("gross_persons", [])

    # 🔴 Spezialpersonen aus Grossreinigung VORAB blockieren
    gross_set = set(x.strip() for x in gross_persons if x)
    blocked_special_names = set(gross_set)

    preblocked_special_ids = set()

    for i, row in df.iterrows():

        name = f"{row['Grad']} {row['Vorname']} {row['Name']}"

        if name in gross_set:

            f = split(row["Funktionen"])

            if excluded(f):
                preblocked_special_ids.add(i)

    persons = []
    special_persons = []

    for i, row in df.iterrows():

        f = split(row["Funktionen"])

        person = {
            "id": i,
            "name": f"{row['Grad']} {row['Vorname']} {row['Name']}",
            "g": split(row["Gruppe"]),
            "f": f
        }

        if excluded(f):
            special_persons.append(person)
        else:
            persons.append(person)

    valid_names = set(
        [p["name"] for p in persons] +
        [p["name"] for p in special_persons]
    )
    name_to_person = {
        p["name"]: p for p in (persons + special_persons)
    }

    pl = Planner(persons + special_persons)
    pl.gross_month = pose_cfg["gross_m"]

    # ❗ KEINE harte Blockierung mehr
    # Spezialpersonen werden über build_plan + assign gesteuert
    pass
    
    # ❗ Spezialpersonen die bereits eingesetzt wurden global merken
    used_special_ids = set()

    folder = os.path.dirname(file)

    out_docx = os.path.join(
        folder,
        f"Probefahrten_{year}.docx"
    )

    out_pdf = os.path.join(
        folder,
        f"Probefahrten_{year}.pdf"
    )

    MAX_TRIES = 20

    for attempt in range(MAX_TRIES):

        pl = Planner(persons + special_persons)
        pl.gross_month = pose_cfg["gross_m"]

        # 🔴 WICHTIG: Blockierungen JEDES MAL neu setzen
        for name in gross_set:
            p = name_to_person.get(name)
            if p and is_special(p):
                pl.h[p["id"]]["count"] = 999

        for sid in preblocked_special_ids:
            pl.h[sid]["count"] = 999

        plan, stats = build_plan(
            pl,
            persons,
            special_persons,
            disabled,
            pose_cfg,
            name_to_person,
            valid_names
        )

        optimize_fairness(plan, persons, stats)
        stats = rebuild_stats(plan, valid_names)

        counts = [d["count"] for d in stats.values()]

        if max(counts) - min(counts) <= 1:
            print(f"✅ Fair solution after {attempt+1}")
            break

        print(f"🔁 Retry {attempt+1}")

    # danach exportieren
    gen_doc(template, out_docx, plan, stats, year)

    # 🔥 DOCX löschen
    if os.path.exists(out_docx):
        os.remove(out_docx)

    with open(out_docx, "rb") as f:
        data = f.read()

    return data
